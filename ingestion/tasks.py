import os
from collections import Counter
import re
import subprocess
import tempfile
import unicodedata
from io import BytesIO
from pathlib import Path

import fitz
from openpyxl import load_workbook
from celery import shared_task
from django.utils import timezone
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from markdownify import markdownify as html_to_markdown

from documents.models import Chunk, Document, ExtractedFact
from providers import embed_texts, generate_chunk_questions
from .models import IngestionJob

QUESTION_GEN_MIN_CHARS = int(os.getenv('QUESTION_GEN_MIN_CHARS', '300'))
QUESTION_GEN_MAX_CHUNKS = int(os.getenv('QUESTION_GEN_MAX_CHUNKS', '120'))
PDF_REPEAT_LINE_THRESHOLD = int(os.getenv('PDF_REPEAT_LINE_THRESHOLD', '12'))
INGEST_MAX_CHUNKS = int(os.getenv('INGEST_MAX_CHUNKS', '1500'))

DOCLING_VENV = os.getenv('DOCLING_VENV_PATH', '/mnt/HC_Volume_105592620/tools/docling/.venv')
DOCLING_PDF_BACKEND = os.getenv('DOCLING_PDF_BACKEND', 'docling_parse')


def repair_suspicious_pdf_tokens(text):
    text = text or ''

    def repair_token(match):
        token = match.group(0)
        letters = sum(ch.isalpha() for ch in token)
        digits = sum(ch.isdigit() for ch in token)
        if letters < 2 or digits == 0:
            return token
        repaired = token
        repaired = re.sub(r'(?<=[A-Za-z])0(?=[A-Za-z])', 't', repaired)
        repaired = re.sub(r'(?<=[A-Za-z])@(?=[A-Za-z])', 'fi', repaired)
        repaired = re.sub(r'(?<=[A-Za-z])§(?=[A-Za-z])', 'ff', repaired)
        repaired = re.sub(r'(?<=[A-Za-z])¢(?=[A-Za-z])', 'ti', repaired)
        return repaired

    text = re.sub(r'\b[\w@§¢]+\b', repair_token, text)
    return text


def remove_repeated_pdf_boilerplate(text, repeat_threshold=PDF_REPEAT_LINE_THRESHOLD):
    lines = [line.strip() for line in (text or '').split('\n')]
    nonempty = [line for line in lines if line]
    counts = Counter(nonempty)

    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned.append('')
            continue
        if counts.get(stripped, 0) >= repeat_threshold and len(stripped) <= 220:
            continue
        if 'Printed for:' in stripped and 'American Payroll Institute' in stripped:
            continue
        cleaned.append(line)

    return '\n'.join(cleaned)


def normalize_extracted_text(text):
    text = repair_suspicious_pdf_tokens((text or '').replace('\x00', ' '))
    text = remove_repeated_pdf_boilerplate(text)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    cleaned_lines = []
    for raw_line in text.split('\n'):
        line = raw_line.strip()
        if not line:
            cleaned_lines.append('')
            continue
        if len(line) <= 3 and any(ch.isdigit() for ch in line):
            continue
        if re.fullmatch(r'[\d\s\|.,:%$\-()]+', line) and sum(ch.isdigit() for ch in line) >= max(4, len(line) // 2):
            continue
        cleaned_lines.append(line)

    paragraphs = []
    current = []
    for line in cleaned_lines:
        if not line:
            if current:
                paragraphs.append(' '.join(current).strip())
                current = []
            continue
        if current and not re.search(r'[.!?:]$|:$', current[-1]) and line[:1].islower():
            current.append(line)
        else:
            current.append(line)
    if current:
        paragraphs.append(' '.join(current).strip())

    normalized = '\n\n'.join(p for p in paragraphs if p)
    normalized = re.sub(r'\n{3,}', '\n\n', normalized)
    return normalized.strip()


EMBEDDING_MAX_INPUT_CHARS = 2000


def build_chunks(text, chunk_size=1000, overlap=None):
    text = normalize_extracted_text(text)
    if not text:
        return []

    if overlap is None:
        overlap = max(80, int(chunk_size * 0.2))

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=['\n\n', '\n', ' ', ''],
        length_function=len,
        is_separator_regex=False,
    )
    chunks = [chunk.strip() for chunk in splitter.split_text(text) if chunk and chunk.strip()]

    deduped = []
    last = None
    for chunk in chunks:
        if chunk != last:
            deduped.append(chunk)
            last = chunk

    if len(deduped) > INGEST_MAX_CHUNKS:
        raise ValueError(
            f'Extraction produced {len(deduped)} chunks, exceeding safety cap of {INGEST_MAX_CHUNKS}. Likely noisy or repetitive document extraction.'
        )

    return deduped


def enforce_embedding_limit(chunks, max_chars=EMBEDDING_MAX_INPUT_CHARS):
    safe_chunks = []
    for chunk in chunks:
        chunk = (chunk or '').strip()
        if not chunk:
            continue
        if len(chunk) <= max_chars:
            safe_chunks.append(chunk)
            continue

        start = 0
        while start < len(chunk):
            end = min(start + max_chars, len(chunk))
            if end < len(chunk):
                split_at = chunk.rfind(' ', start, end)
                if split_at > start + 100:
                    end = split_at
            piece = chunk[start:end].strip()
            if piece:
                safe_chunks.append(piece)
            if end >= len(chunk):
                break
            start = end
    return safe_chunks


def extract_pdf_text(raw):
    def _merge_pdf_blocks(page):
        blocks = page.get_text('blocks') or []
        normalized = []
        for block in blocks:
            if len(block) < 5:
                continue
            x0, y0, x1, y1, text = block[:5]
            text = re.sub(r'\s+', ' ', (text or '')).strip()
            if not text:
                continue
            normalized.append({
                'x0': float(x0),
                'y0': float(y0),
                'x1': float(x1),
                'y1': float(y1),
                'text': text,
            })

        normalized.sort(key=lambda item: (round(item['y0'], 1), item['x0']))
        merged_lines = []
        row_tolerance = 6.0
        for block in normalized:
            placed = False
            for row in merged_lines:
                if abs(row['y0'] - block['y0']) <= row_tolerance:
                    row['blocks'].append(block)
                    row['y0'] = min(row['y0'], block['y0'])
                    placed = True
                    break
            if not placed:
                merged_lines.append({'y0': block['y0'], 'blocks': [block]})

        line_texts = []
        for row in sorted(merged_lines, key=lambda item: item['y0']):
            row_blocks = sorted(row['blocks'], key=lambda item: item['x0'])
            parts = []
            for block in row_blocks:
                token = block['text']
                if parts and len(token) <= 4 and re.fullmatch(r'[12][A-Z]|[A-Z0-9]{1,4}', token):
                    parts[-1] = f"{parts[-1]} {token}"
                else:
                    parts.append(token)
            line = ' '.join(parts).strip()
            if line:
                line_texts.append(line)
        return '\n'.join(line_texts)

    pdf = fitz.open(stream=raw, filetype='pdf')
    try:
        pages = []
        for page in pdf:
            block_text = _merge_pdf_blocks(page)
            fallback_text = page.get_text('text') or ''
            chosen = block_text.strip() if len(block_text.strip()) >= len(fallback_text.strip()) * 0.7 else fallback_text.strip()
            if chosen:
                pages.append(chosen)
    finally:
        pdf.close()
    return '\n\n'.join(page.strip() for page in pages if page and page.strip())


def extract_pdf_text_docling(document):
    if not document.file:
        return ''
    python_path = Path(DOCLING_VENV) / 'bin' / 'python'
    if not python_path.exists():
        raise RuntimeError(f'Docling venv not found at {python_path}')

    suffix = Path(document.filename or 'document.pdf').suffix or '.pdf'
    with document.file.open('rb') as fh, tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(fh.read())
        source_path = Path(tmp.name)

    cli_path = Path(DOCLING_VENV) / 'bin' / 'docling'
    if not cli_path.exists():
        raise RuntimeError(f'Docling CLI not found at {cli_path}')
    output_dir = Path(tempfile.mkdtemp(prefix='docling-out-'))
    try:
        try:
            completed = subprocess.run(
                [
                    str(cli_path),
                    '--from', 'pdf',
                    '--to', 'md',
                    '--output', str(output_dir),
                    '--pdf-backend', DOCLING_PDF_BACKEND,
                    str(source_path),
                ],
                capture_output=True,
                text=True,
                check=True,
                env={**os.environ, 'DOCLING_PDF_BACKEND': DOCLING_PDF_BACKEND},
            )
            md_files = sorted(output_dir.glob('*.md'))
            if not md_files:
                raise RuntimeError(completed.stderr.strip() or completed.stdout.strip() or 'Docling produced no markdown output')
            return md_files[0].read_text().strip()
        except subprocess.CalledProcessError as exc:
            raise RuntimeError(exc.stderr.strip() or exc.stdout.strip() or 'Docling extraction failed')
    finally:
        for path in output_dir.glob('*'):
            try:
                path.unlink()
            except Exception:
                pass
        try:
            output_dir.rmdir()
        except Exception:
            pass
        try:
            source_path.unlink(missing_ok=True)
        except Exception:
            pass


def extract_docx_text(raw):
    doc = DocxDocument(BytesIO(raw))
    parts = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or '').strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            values = [(cell.text or '').strip() for cell in row.cells]
            values = [value for value in values if value]
            if values:
                parts.append(' | '.join(values))
    return '\n\n'.join(parts)


def extract_text_document(raw):
    return raw.decode('utf-8', errors='ignore')


def extract_xlsx_text(raw, filename='workbook.xlsx'):
    def normalize_value(value):
        if value is None:
            return ''
        return str(value).strip()

    def row_has_content(values):
        return any(v != '' for v in values)

    sections = [f'Workbook: {filename}']
    workbook = load_workbook(BytesIO(raw), data_only=True, read_only=True)

    for sheet in workbook.worksheets:
        sections.append('')
        sections.append(f'Sheet: {sheet.title}')

        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            sections.append('(empty sheet)')
            continue

        header_row = rows[0]
        columns = [normalize_value(col) or f'column_{idx + 1}' for idx, col in enumerate(header_row)]
        sections.append('')
        sections.append('Columns:')
        sections.append(' | '.join(columns))

        row_number = 0
        for row in rows[1:]:
            values = [normalize_value(cell) for cell in row]
            if len(values) < len(columns):
                values.extend([''] * (len(columns) - len(values)))
            elif len(values) > len(columns):
                extra_count = len(values) - len(columns)
                columns_extended = columns + [f'column_{len(columns) + i + 1}' for i in range(extra_count)]
            else:
                columns_extended = columns

            if len(values) <= len(columns):
                columns_extended = columns

            if not row_has_content(values):
                continue

            row_number += 1
            sections.append('')
            sections.append(f'Row {row_number}')
            for col_name, value in zip(columns_extended, values):
                sections.append(f'{col_name}: {value}')

        if row_number == 0:
            sections.append('')
            sections.append('(no non-empty rows)')

    return '\n'.join(sections).strip()


def extract_html_text(raw):
    html = raw.decode('utf-8', errors='ignore')
    return html_to_markdown(html)


def extract_document_text(document, version, extractor=IngestionJob.EXTRACTOR_STANDARD):
    if not document.file:
        return (version.extraction_metadata_json or {}).get('raw_text', '')

    with document.file.open('rb') as fh:
        raw = fh.read()

    filename = (document.filename or '').lower()
    mime_type = (document.mime_type or '').lower()

    if filename.endswith('.pdf') or mime_type == 'application/pdf':
        if extractor == IngestionJob.EXTRACTOR_DOCLING:
            return extract_pdf_text_docling(document)
        return extract_pdf_text(raw)
    if filename.endswith('.docx') or mime_type in {
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword',
    }:
        return extract_docx_text(raw)
    if filename.endswith('.xlsx') or mime_type in {
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel',
    }:
        return extract_xlsx_text(raw, filename=document.filename or 'workbook.xlsx')
    if filename.endswith('.md') or mime_type in {'text/markdown', 'text/x-markdown'}:
        return extract_text_document(raw)
    if filename.endswith('.txt') or mime_type.startswith('text/plain'):
        return extract_text_document(raw)
    if filename.endswith('.html') or filename.endswith('.htm') or mime_type == 'text/html':
        return extract_html_text(raw)

    return extract_text_document(raw)


def normalize_structure_text(text):
    text = unicodedata.normalize('NFKC', text or '')
    text = text.replace('’', "'").replace('“', '"').replace('”', '"')
    text = text.replace('–', '-').replace('—', '-')
    text = text.replace('Ɵ', 'Ti').replace('Ō', 'O').replace('ﬁ', 'fi').replace('ﬂ', 'fl')
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def is_heading_candidate(text):
    text = normalize_structure_text(text).strip(' •\t-')
    if not text:
        return False
    if len(text) < 3 or len(text) > 70:
        return False
    if ':' in text:
        return False
    lowered = text.lower()
    if any(term in lowered for term in ['president', 'chief executive', 'kevin clayton']):
        return False
    words = text.split()
    if len(words) >= 3 and all(word[:1].isupper() for word in words if word[:1].isalpha()):
        policy_terms = ['holiday', 'pto', 'leave', 'benefit', 'conduct', 'policy', 'harassment', 'eligibility', 'compensation', 'attendance']
        if not any(term in lowered for term in policy_terms):
            return False
    return text == text.title() or bool(re.match(r'^\d+[\-.]\s+', text))


def analyze_chunk_structure(chunk_text):
    text = normalize_structure_text((chunk_text or '').strip())
    raw_lines = [line.rstrip() for line in text.split('\n') if line.strip()]
    cleaned_lines = [line.strip() for line in raw_lines]
    heading_candidates = []
    list_lines = []

    inline_heading_patterns = [
        r'^(\d+(?:[\-.]\d+)*)\.\s+([A-Z][A-Za-z’\'&/\- ]{2,60}?)(?=\s+[A-Z][a-z]|\s+Full-|\s+If\s|\s+All\s|\s+Employees\s|:|$)',
        r'^(Section\s+\d+\s+[–-]\s+[A-Z][^.:]{2,80})',
    ]

    for line in cleaned_lines[:8]:
        normalized = re.sub(r'\s+', ' ', line).strip(' •\t-')
        if not normalized:
            continue
        if is_heading_candidate(normalized):
            heading_candidates.append(normalized)
        else:
            for pattern in inline_heading_patterns:
                match = re.match(pattern, normalized)
                if match:
                    heading_candidates.append(match.group(0).strip())
                    break

        if re.match(r'^(?:[\-•*]|\d+[.)])\s+', line):
            list_lines.append(re.sub(r'^(?:[\-•*]|\d+[.)])\s+', '', normalized).strip())

        if '•' in line and line.count('•') >= 2:
            parts = [part.strip(' •\t-') for part in line.split('•') if part.strip(' •\t-')]
            if len(parts) > 1:
                for item in parts[1:]:
                    cleaned_item = re.sub(r'\s+', ' ', item)
                    cleaned_item = re.split(r'(?=\b(?:To be|If |Unless |Part-|Employees |Holiday pay|Time missed)\b)', cleaned_item, maxsplit=1)[0].strip(' •\t-')
                    if len(cleaned_item) >= 3:
                        list_lines.append(cleaned_item)

    deduped_headings = []
    seen_headings = set()
    for heading in heading_candidates:
        if heading not in seen_headings:
            deduped_headings.append(heading)
            seen_headings.add(heading)

    deduped_list_lines = []
    seen_items = set()
    for item in list_lines:
        normalized_item = re.sub(r'\s+', ' ', item).strip(' •\t-')
        if not normalized_item or normalized_item in seen_items:
            continue
        seen_items.add(normalized_item)
        deduped_list_lines.append(normalized_item)

    dominant_heading = deduped_headings[0] if deduped_headings else ''
    return {
        'heading_candidates': deduped_headings,
        'dominant_heading': dominant_heading,
        'list_lines': deduped_list_lines,
        'list_count': len(deduped_list_lines),
        'line_count': len(cleaned_lines),
        'has_heading': bool(deduped_headings),
        'has_list': bool(deduped_list_lines),
    }


def should_generate_llm_questions(chunk_text, structure=None):
    structure = structure or analyze_chunk_structure(chunk_text)
    text = (chunk_text or '').strip()
    if len(text) < QUESTION_GEN_MIN_CHARS:
        return False
    lowered = text.lower()
    if 'contents' in lowered and '........' in lowered:
        return False
    if lowered.count('�') >= 3:
        return False
    if re.fullmatch(r'[\W\d_\s]+', text):
        return False
    useful_terms = ['holiday', 'pto', 'leave', 'benefit', 'eligib', 'coverage', 'harassment', 'complaint', 'report', 'conduct', 'policy', 'handbook', 'at-will']
    if structure.get('has_list'):
        return True
    return any(term in lowered for term in useful_terms)


def infer_chunk_questions(document, chunk_text, structure=None, use_llm=False):
    structure = structure or analyze_chunk_structure(chunk_text)
    heading = structure.get('dominant_heading', '')
    lowered = (chunk_text or '').lower()
    questions = []

    if 'holiday' in lowered:
        questions.extend([
            'What paid holidays does the company observe?',
            'What holidays are employees off?',
            'Which holidays are recognized by the company?',
            'Is the day after Thanksgiving a paid holiday?',
        ])
    if any(token in lowered for token in ['pto', 'vacation', 'paid time off']):
        questions.extend([
            'How does PTO work?',
            'How much paid time off do employees get?',
            'What are the vacation or PTO rules?',
        ])
    if any(token in lowered for token in ['eligib', 'coverage', 'benefit']):
        questions.extend([
            'When are employees eligible for benefits?',
            'What are the benefits eligibility rules?',
            'When does employee coverage begin?',
        ])
    if any(token in lowered for token in ['harassment', 'complaint', 'report', 'prohibited conduct']):
        questions.extend([
            'How should employees report harassment?',
            'Who should prohibited conduct be reported to?',
            'What happens after a harassment complaint is made?',
        ])
    if any(token in lowered for token in ['at-will', 'employment at will', 'contract', 'handbook']):
        questions.extend([
            'Does the handbook create an employment contract?',
            'Is employment at will?',
            'What is the purpose of this handbook?',
        ])
    if structure.get('has_list'):
        questions.append('What items are listed in this policy section?')
    if heading and is_heading_candidate(heading):
        questions.append(f'What does the {heading} policy cover?')

    if use_llm and should_generate_llm_questions(chunk_text, structure=structure):
        llm_questions = generate_chunk_questions(chunk_text)
        if llm_questions:
            questions = llm_questions

    if not questions:
        questions.extend([
            f'What policy or rules are described in {document.filename}?',
            'What employee guidance is given in this section?',
        ])

    deduped = []
    seen = set()
    for question in questions:
        normalized = question.strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(normalized)
    return '\n'.join(deduped[:5])


def build_chunk_metadata_text(document, chunk_text, structure=None):
    structure = structure or analyze_chunk_structure(chunk_text)
    heading = structure.get('dominant_heading', '')
    list_lines = structure.get('list_lines', [])[:5]
    list_preview = '; '.join(list_lines)
    keywords = []
    lowered = (chunk_text or '').lower()
    for token in ['holiday', 'holidays', 'paid', 'pto', 'leave', 'eligibility', 'benefits', 'vacation', 'coverage', 'schedule']:
        if token in lowered:
            keywords.append(token)
    parts = [
        f'Document: {document.filename}',
        f'Heading: {heading or "(none)"}',
        f'Chunk type: {"list" if structure.get("has_list") else "policy text"}',
        f'List count: {structure.get("list_count", 0)}',
    ]
    if keywords:
        parts.append(f'Keywords: {", ".join(keywords)}')
    if list_preview:
        parts.append(f'List preview: {list_preview}')
    parts.append(f'Chunk summary: {(chunk_text or "")[:400].replace(chr(10), " ")}')
    return '\n'.join(parts).strip()


def extract_chunk_facts(chunk_text, structure=None):
    facts = []
    text = (chunk_text or '').strip()
    if not text:
        return facts

    structure = structure or analyze_chunk_structure(chunk_text)
    lines = [line.strip(' •\t-') for line in text.split('\n') if line.strip()]
    for line in lines:
        normalized = re.sub(r'\s+', ' ', line).strip()
        if len(normalized) < 4:
            continue
        if normalized in structure.get('heading_candidates', []):
            facts.append({
                'fact_type': ExtractedFact.FACT_HEADING,
                'label': normalized,
                'value_text': normalized,
                'normalized_text': normalized.lower(),
                'confidence': 0.75,
                'metadata_json': {'pattern': 'heading_like', 'dominant_heading': structure.get('dominant_heading', '')},
            })
        if re.match(r'^(?:[\-•*]|\d+[.)])\s+', line):
            facts.append({
                'fact_type': ExtractedFact.FACT_LIST_ITEM,
                'label': structure.get('dominant_heading', '')[:255],
                'value_text': normalized,
                'normalized_text': normalized.lower(),
                'confidence': 0.9,
                'metadata_json': {
                    'pattern': 'bullet_or_numbered_line',
                    'dominant_heading': structure.get('dominant_heading', ''),
                    'list_count': structure.get('list_count', 0),
                },
            })
        elif ':' in normalized and len(normalized) <= 240:
            label, value = [part.strip() for part in normalized.split(':', 1)]
            if label and value:
                facts.append({
                    'fact_type': ExtractedFact.FACT_POLICY,
                    'label': label[:255],
                    'value_text': value,
                    'normalized_text': f'{label} {value}'.lower(),
                    'confidence': 0.7,
                    'metadata_json': {
                        'pattern': 'label_value',
                        'dominant_heading': structure.get('dominant_heading', ''),
                    },
                })

    sentences = re.split(r'(?<=[.!?])\s+', text)
    for sentence in sentences:
        sentence = re.sub(r'\s+', ' ', sentence).strip()
        if len(sentence) < 30 or len(sentence) > 320:
            continue
        lowered = sentence.lower()
        if any(keyword in lowered for keyword in ['holiday', 'pto', 'leave', 'eligible', 'coverage', 'benefit', 'schedule', 'vacation']):
            facts.append({
                'fact_type': ExtractedFact.FACT_POLICY,
                'label': structure.get('dominant_heading', '')[:255],
                'value_text': sentence,
                'normalized_text': lowered,
                'confidence': 0.6,
                'metadata_json': {
                    'pattern': 'policy_sentence',
                    'dominant_heading': structure.get('dominant_heading', ''),
                },
            })

    deduped = []
    seen = set()
    for fact in facts:
        key = (fact['fact_type'], fact['label'], fact['normalized_text'])
        if key in seen:
            continue
        seen.add(key)
        deduped.append(fact)
    return deduped


@shared_task(bind=True, autoretry_for=(Exception,), retry_backoff=True, retry_kwargs={'max_retries': 1})
def ingest_document_task(self, ingestion_job_id):
    job = IngestionJob.objects.select_related('document', 'document_version', 'tenant', 'workspace').get(id=ingestion_job_id)
    document = job.document
    version = job.document_version

    job.status = IngestionJob.STATUS_RUNNING
    job.stage = 'extracting'
    job.started_at = timezone.now()
    job.error_text = ''
    job.save(update_fields=['status', 'stage', 'started_at', 'error_text'])

    document.status = Document.STATUS_PROCESSING
    document.save(update_fields=['status', 'updated_at'])

    try:
        extracted_text = extract_document_text(document, version, extractor=job.extractor)
        cleaned_text = normalize_extracted_text(extracted_text)

        job.stage = 'embedding'
        job.save(update_fields=['stage'])

        chunk_size = job.workspace.default_chunk_size or 1000
        chunks = build_chunks(cleaned_text, chunk_size=chunk_size, overlap=max(80, int(chunk_size * 0.2)))
        chunks = enforce_embedding_limit(chunks)
        metadata_texts = []
        question_texts = []
        chunk_structures = []
        last_heading_for_metadata = ''
        llm_question_budget = QUESTION_GEN_MAX_CHUNKS
        for chunk_text in chunks:
            structure = analyze_chunk_structure(chunk_text)
            dominant_heading = structure.get('dominant_heading') or last_heading_for_metadata
            if dominant_heading:
                structure['dominant_heading'] = dominant_heading
                last_heading_for_metadata = dominant_heading
            chunk_structures.append(structure)
            metadata_texts.append(build_chunk_metadata_text(document, chunk_text, structure=structure))
            use_llm = llm_question_budget > 0 and should_generate_llm_questions(chunk_text, structure=structure)
            question_text = infer_chunk_questions(document, chunk_text, structure=structure, use_llm=use_llm)
            if use_llm:
                llm_question_budget -= 1
            question_texts.append(question_text)
        vectors = embed_texts(chunks) if chunks else []
        metadata_vectors = embed_texts(metadata_texts) if metadata_texts else []
        question_vectors = embed_texts(question_texts) if question_texts else []
        Chunk.objects.filter(document_version=version).delete()
        ExtractedFact.objects.filter(document_version=version).delete()
        created_chunks = []
        last_heading = ''
        for idx, chunk_text in enumerate(chunks):
            structure = chunk_structures[idx] if idx < len(chunk_structures) else analyze_chunk_structure(chunk_text)
            dominant_heading = structure.get('dominant_heading') or last_heading
            if dominant_heading:
                structure['dominant_heading'] = dominant_heading
                last_heading = dominant_heading

            metadata_json = {
                'stub': False,
                'dominant_heading': dominant_heading,
                'heading_candidates': structure.get('heading_candidates', []),
                'list_count': structure.get('list_count', 0),
                'has_list': structure.get('has_list', False),
                'line_count': structure.get('line_count', 0),
            }
            lowered_chunk = normalize_structure_text(chunk_text).lower()
            if 'holiday' in lowered_chunk:
                metadata_json['debug_structure'] = {
                    'matched_debug_rule': True,
                    'raw_preview': chunk_text[:1200],
                    'analyze_result': {
                        'dominant_heading': structure.get('dominant_heading', ''),
                        'heading_candidates': structure.get('heading_candidates', []),
                        'list_count': structure.get('list_count', 0),
                        'has_list': structure.get('has_list', False),
                        'list_lines_preview': (structure.get('list_lines', []) or [])[:10],
                    },
                }

            chunk = Chunk.objects.create(
                tenant=job.tenant,
                workspace=job.workspace,
                document=document,
                document_version=version,
                chunk_index=idx,
                text=chunk_text,
                metadata_text=metadata_texts[idx] if idx < len(metadata_texts) else '',
                question_text=question_texts[idx] if idx < len(question_texts) else '',
                token_count=max(1, len(chunk_text) // 4),
                metadata_json=metadata_json,
                embedding=vectors[idx] if idx < len(vectors) else None,
                metadata_embedding=metadata_vectors[idx] if idx < len(metadata_vectors) else None,
                question_embedding=question_vectors[idx] if idx < len(question_vectors) else None,
            )
            created_chunks.append(chunk)

            for fact in extract_chunk_facts(chunk_text, structure=structure):
                ExtractedFact.objects.create(
                    tenant=job.tenant,
                    workspace=job.workspace,
                    document=document,
                    document_version=version,
                    chunk=chunk,
                    fact_type=fact['fact_type'],
                    label=fact.get('label', '')[:255],
                    value_text=fact['value_text'],
                    normalized_text=fact.get('normalized_text', ''),
                    metadata_json=fact.get('metadata_json', {}),
                    confidence=fact.get('confidence', 0.0),
                )

        version.parse_status = 'ready'
        version.extraction_metadata_json = {
            **(version.extraction_metadata_json or {}),
            'raw_text': cleaned_text,
            'raw_text_preview': cleaned_text[:500],
            'chunk_count': len(chunks),
            'fact_count': ExtractedFact.objects.filter(document_version=version).count(),
            'extractor': job.extractor,
        }
        version.save(update_fields=['parse_status', 'extraction_metadata_json'])

        document.status = Document.STATUS_READY
        document.save(update_fields=['status', 'updated_at'])

        job.status = IngestionJob.STATUS_SUCCEEDED
        job.stage = 'done'
        job.finished_at = timezone.now()
        job.save(update_fields=['status', 'stage', 'finished_at'])
        return {'chunk_count': len(chunks)}
    except Exception as exc:
        job.status = IngestionJob.STATUS_FAILED
        job.stage = 'failed'
        job.error_text = str(exc)
        job.finished_at = timezone.now()
        job.save(update_fields=['status', 'stage', 'error_text', 'finished_at'])
        document.status = Document.STATUS_FAILED
        document.save(update_fields=['status', 'updated_at'])
        raise
